"""
ExonEditor Backend — FastAPI
Hospedagem: Render
Funcionalidades:
  - Busca NCBI GenBank
  - Parsing Biopython (éxons, proteínas, região sugerida)
  - Geração de documento Word (.docx)
  - Análise de códons e stop codon
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
import urllib.request, urllib.parse, json, re, time, io
from Bio import SeqIO
from docx import Document
from docx.shared import RGBColor, Pt

app = FastAPI(title="ExonEditor API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Código Genético ──────────────────────────────────────────
CODIGO_GENETICO = {
    'TTT':'Phe','TTC':'Phe','TTA':'Leu','TTG':'Leu',
    'CTT':'Leu','CTC':'Leu','CTA':'Leu','CTG':'Leu',
    'ATT':'Ile','ATC':'Ile','ATA':'Ile','ATG':'Met',
    'GTT':'Val','GTC':'Val','GTA':'Val','GTG':'Val',
    'TCT':'Ser','TCC':'Ser','TCA':'Ser','TCG':'Ser',
    'CCT':'Pro','CCC':'Pro','CCA':'Pro','CCG':'Pro',
    'ACT':'Thr','ACC':'Thr','ACA':'Thr','ACG':'Thr',
    'GCT':'Ala','GCC':'Ala','GCA':'Ala','GCG':'Ala',
    'TAT':'Tyr','TAC':'Tyr','TAA':'Stop','TAG':'Stop',
    'CAT':'His','CAC':'His','CAA':'Gln','CAG':'Gln',
    'AAT':'Asn','AAC':'Asn','AAA':'Lys','AAG':'Lys',
    'GAT':'Asp','GAC':'Asp','GAA':'Glu','GAG':'Glu',
    'TGT':'Cys','TGC':'Cys','TGA':'Stop','TGG':'Trp',
    'CGT':'Arg','CGC':'Arg','CGA':'Arg','CGG':'Arg',
    'AGT':'Ser','AGC':'Ser','AGA':'Arg','AGG':'Arg',
    'GGT':'Gly','GGC':'Gly','GGA':'Gly','GGG':'Gly',
}

AA_1_TO_3 = {
    'A':'Ala','R':'Arg','N':'Asn','D':'Asp','C':'Cys','Q':'Gln','E':'Glu',
    'G':'Gly','H':'His','I':'Ile','L':'Leu','K':'Lys','M':'Met','F':'Phe',
    'P':'Pro','S':'Ser','T':'Thr','W':'Trp','Y':'Tyr','V':'Val','*':'Stop',
}

AA_SINONIMOS = {
    'A':'Ala','R':'Arg','N':'Asn','D':'Asp','C':'Cys','Q':'Gln','E':'Glu',
    'G':'Gly','H':'His','I':'Ile','L':'Leu','K':'Lys','M':'Met','F':'Phe',
    'P':'Pro','S':'Ser','T':'Thr','W':'Trp','Y':'Tyr','V':'Val',
    'ALA':'Ala','ARG':'Arg','ASN':'Asn','ASP':'Asp','CYS':'Cys','GLN':'Gln',
    'GLU':'Glu','GLY':'Gly','HIS':'His','ILE':'Ile','LEU':'Leu','LYS':'Lys',
    'MET':'Met','PHE':'Phe','PRO':'Pro','SER':'Ser','THR':'Thr','TRP':'Trp',
    'TYR':'Tyr','VAL':'Val',
}

STOP_CODONS = {'TAA', 'TAG', 'TGA'}

ENTREZ = 'https://eutils.ncbi.nlm.nih.gov/entrez/eutils'
EMAIL  = 'exoneditor@app.tool'


# ── NCBI helpers ────────────────────────────────────────────
def _get(url: str, timeout: int = 30) -> str:
    for attempt in range(3):
        try:
            with urllib.request.urlopen(url, timeout=timeout) as r:
                return r.read().decode('utf-8', errors='replace')
        except Exception as e:
            if attempt == 2:
                raise RuntimeError(f"Erro NCBI: {e}")
            time.sleep(1.5)


def ncbi_buscar(termo: str, email: str = EMAIL, retmax: int = 8) -> list:
    p = urllib.parse.urlencode({
        'db': 'nuccore', 'term': termo, 'retmax': retmax,
        'retmode': 'json', 'email': email, 'tool': 'ExonEditor'
    })
    dados = json.loads(_get(f"{ENTREZ}/esearch.fcgi?{p}"))
    ids = dados.get('esearchresult', {}).get('idlist', [])
    if not ids:
        return []
    p2 = urllib.parse.urlencode({
        'db': 'nuccore', 'id': ','.join(ids),
        'retmode': 'json', 'email': email, 'tool': 'ExonEditor'
    })
    summ = json.loads(_get(f"{ENTREZ}/esummary.fcgi?{p2}"))
    result = summ.get('result', {})
    uids = result.get('uids', ids)
    regs = []
    for uid in uids:
        info = result.get(uid, {})
        regs.append({
            'id': uid,
            'acc': info.get('accessionversion', uid),
            'titulo': info.get('title', '—'),
            'len': info.get('slen', '?'),
        })
    return regs


def ncbi_fetch_genbank(acc: str, email: str = EMAIL,
                       seq_start: int = None, seq_stop: int = None) -> str:
    params = {
        'db': 'nuccore', 'id': acc, 'rettype': 'gb',
        'retmode': 'text', 'email': email, 'tool': 'ExonEditor'
    }
    if seq_start and seq_stop:
        params['seq_start'] = seq_start
        params['seq_stop']  = seq_stop
        params['strand']    = 1
    p = urllib.parse.urlencode(params)
    return _get(f"{ENTREZ}/efetch.fcgi?{p}", timeout=45)


def parsear_genbank(gb: str) -> dict:
    """Parser usando Biopython."""
    from io import StringIO
    record = SeqIO.read(StringIO(gb), 'genbank')
    seq = str(record.seq).upper()

    res = {
        'sequencia': seq,
        'exons': [],
        'proteinas': [],
        'organismo': '',
        'definicao': record.description,
        'accession': record.id,
        'regiao_sugerida': None,
        'proteina_nascente': '',
        'proteina_madura': [],
    }

    for f in record.features:
        if f.type == 'source':
            res['organismo'] = f.qualifiers.get('organism', [''])[0]
            break

    for f in record.features:
        if f.type == 'exon':
            nums = f.qualifiers.get('number', [None])
            num  = int(nums[0]) if nums[0] else len(res['exons']) + 1
            res['exons'].append({
                'numero': num,
                'inicio': int(f.location.start) + 1,
                'fim':    int(f.location.end),
            })
    res['exons'].sort(key=lambda x: x['numero'])

    for f in record.features:
        if f.type == 'CDS':
            trans = f.qualifiers.get('translation', [''])[0]
            prod  = f.qualifiers.get('product', ['—'])[0]
            pid   = f.qualifiers.get('protein_id', ['—'])[0]
            try:
                coords_cds = [(int(p.start) + 1, int(p.end))
                              for p in f.location.parts]
            except Exception:
                coords_cds = []
            res['proteinas'].append({
                'produto':    prod,
                'protein_id': pid,
                'sequencia':  trans,
                'n_aa':       len(trans),
                'coords_cds': coords_cds,
            })
            if trans and not res['proteina_nascente']:
                res['proteina_nascente'] = trans

    for f in record.features:
        if f.type == 'mat_peptide':
            try:
                coords = [(int(p.start) + 1, int(p.end))
                          for p in f.location.parts]
            except Exception:
                coords = [(int(f.location.start) + 1, int(f.location.end))]
            res['proteina_madura'].append({
                'produto': f.qualifiers.get('product', ['—'])[0],
                'coords':  coords,
            })

    if res['exons']:
        res['regiao_sugerida'] = {
            'de':  res['exons'][0]['inicio'],
            'ate': res['exons'][-1]['fim'],
        }

    # Proteína madura por offset
    if res['proteinas'] and res['proteina_madura'] and res['proteina_nascente']:
        coords_cds = res['proteinas'][0].get('coords_cds', [])
        coords_mat = res['proteina_madura'][0].get('coords', [])
        if coords_cds and coords_mat:
            mat_ini = coords_mat[0][0]
            offset_nt = 0
            for ini_c, fim_c in coords_cds:
                if mat_ini > fim_c:
                    offset_nt += fim_c - ini_c + 1
                elif mat_ini >= ini_c:
                    offset_nt += mat_ini - ini_c
                    break
            offset_aa = offset_nt // 3
            nasc = res['proteina_nascente']
            res['proteina_madura_seq'] = nasc[offset_aa:] if offset_aa < len(nasc) else ''
        else:
            nasc = res['proteina_nascente']
            res['proteina_madura_seq'] = nasc[24:] if len(nasc) > 24 else ''
    else:
        res['proteina_madura_seq'] = ''

    return res


# ── Geração de documentos ────────────────────────────────────
def cor_hex_rgb(h: str):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def aplicar_run(run, negrito, fonte, tam, cor):
    r, g, b = cor_hex_rgb(cor)
    run.bold = negrito
    run.font.name = fonte
    run.font.size = Pt(tam)
    run.font.color.rgb = RGBColor(r, g, b)

def doc_base():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Pt(36)
    s.left_margin = s.right_margin = Pt(54)
    return doc

def gerar_completo_bytes(seq, cfg, exons, caixa='minuscula',
                          codons_grifar=None, utr_regiao=None, chars=60):
    seq = seq.lower() if caixa == 'minuscula' else seq.upper()
    n = len(seq)
    anot = [None] * n
    for ex in exons:
        for i in range(max(0, ex['inicio']-1), min(n, ex['fim'])):
            anot[i] = ex

    utr_map = [False] * n
    if utr_regiao:
        for i in range(utr_regiao['pos_ini'], min(n, utr_regiao['pos_fim'])):
            utr_map[i] = True

    grifo = [None] * n
    if codons_grifar:
        for g in codons_grifar:
            for i in range(g['pos_ini'], min(n, g['pos_fim'])):
                grifo[i] = g

    doc = doc_base()
    for li in range(0, n, chars):
        tr  = seq[li:li+chars]
        al  = anot[li:li+chars]
        gr  = grifo[li:li+chars]
        utr = utr_map[li:li+chars]
        p   = doc.add_paragraph()
        p.paragraph_format.space_after = p.paragraph_format.space_before = Pt(0)
        p.style.font.name = cfg['fonte']
        i = 0
        while i < len(tr):
            cur_a = al[i]; cur_g = gr[i]; cur_u = utr[i]
            j = i + 1
            while j < len(tr) and al[j] is cur_a and gr[j] is cur_g and utr[j] == cur_u:
                j += 1
            run = p.add_run(tr[i:j])
            if cur_g is not None:
                aplicar_run(run, True, cfg['fonte'], cfg['tamanho'], cur_g['cor'])
            elif cur_u:
                aplicar_run(run, True, cfg['fonte'], cfg['tamanho'],
                            utr_regiao.get('cor', '#6699cc'))
            elif cur_a is None:
                aplicar_run(run, False, cfg['fonte'], cfg['tamanho'], cfg['cor'])
            else:
                aplicar_run(run, True, cur_a['fonte'], cur_a['tamanho'], cur_a['cor'])
            i = j

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def gerar_so_exons_bytes(seq, exons, caixa='minuscula',
                          codons_grifar=None, utr_regiao=None, chars=60):
    seq_orig = seq.lower() if caixa == 'minuscula' else seq.upper()
    exons_ord = sorted(exons, key=lambda x: x['inicio'])
    seq_ex = ''; mapa_pos = []; mapa_fmt = []
    for idx, ex in enumerate(exons_ord):
        cor = '#000000' if idx % 2 == 0 else '#e67e00'
        obj = {'cor': cor, 'fonte': ex.get('fonte', 'Courier New'),
               'tamanho': ex.get('tamanho', 11)}
        ini = max(0, ex['inicio']-1); fim = min(len(seq_orig), ex['fim'])
        for p in range(ini, fim):
            seq_ex += seq_orig[p]; mapa_pos.append(p); mapa_fmt.append(obj)

    grifo_map = {}
    if codons_grifar:
        for g in codons_grifar:
            for p in range(g['pos_ini'], g['pos_fim']):
                grifo_map[p] = g
    utr_set = set()
    if utr_regiao:
        for p in range(utr_regiao['pos_ini'], utr_regiao['pos_fim']):
            utr_set.add(p)

    doc = doc_base()
    for li in range(0, len(seq_ex), chars):
        fatia = seq_ex[li:li+chars]
        pos_l = mapa_pos[li:li+chars]
        fmt_l = mapa_fmt[li:li+chars]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = p.paragraph_format.space_before = Pt(0)
        i = 0
        while i < len(fatia):
            orig_p  = pos_l[i]; cur_fmt = fmt_l[i]
            cur_g   = grifo_map.get(orig_p); cur_u = orig_p in utr_set
            j = i + 1
            while j < len(fatia):
                if (fmt_l[j] is cur_fmt and grifo_map.get(pos_l[j]) is cur_g
                        and (pos_l[j] in utr_set) == cur_u):
                    j += 1
                else:
                    break
            run = p.add_run(fatia[i:j])
            if cur_g:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'], cur_g['cor'])
            elif cur_u and utr_regiao:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'],
                            utr_regiao.get('cor', '#6699cc'))
            else:
                aplicar_run(run, True, cur_fmt['fonte'], cur_fmt['tamanho'], cur_fmt['cor'])
            i = j

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Análise de códons ────────────────────────────────────────
def construir_mapa_cds(seq, coords_cds):
    mapa = []
    for ini, fim in coords_cds:
        for p in range(max(1, ini), min(len(seq)+1, fim+1)):
            mapa.append(p)
    return mapa

def analisar_primeiros_codons(seq, proteina_nascente, coords_cds, aas_entrada):
    aas = []
    for a in aas_entrada:
        norm = AA_SINONIMOS.get(a.strip().upper())
        if not norm:
            return {'erro': f"Aminoácido não reconhecido: '{a}'"}
        aas.append(norm)
    if len(aas) < 3:
        return {'erro': 'Mínimo 3 aminoácidos.'}
    n = len(aas)
    prot_3 = [AA_1_TO_3.get(c.upper(), '?') for c in proteina_nascente[:n]]
    matches = [aas[i] == prot_3[i] for i in range(n)]
    mapa = construir_mapa_cds(seq, coords_cds) if coords_cds else []
    pos_gene = []
    codons_reais = []
    for i in range(n):
        idx = i * 3
        if idx < len(mapa):
            p = mapa[idx]
            pos_gene.append(p)
            codon = seq[p-1:p+2].upper() if p+2 <= len(seq) else '???'
            codons_reais.append(codon)
        else:
            pos_gene.append(0); codons_reais.append('???')
    return {
        'erro': None,
        'encontrado': all(matches),
        'aas_fornecidos': aas,
        'aas_na_proteina': prot_3,
        'codons_reais': codons_reais,
        'posicoes_gene': pos_gene,
        'matches': matches,
    }

def localizar_codon_num(seq, coords_cds, numero_aa):
    mapa = construir_mapa_cds(seq, coords_cds)
    i0 = (numero_aa - 1) * 3
    i1 = numero_aa * 3
    if i1 > len(mapa):
        return {'erro': f'Aminoácido {numero_aa} além da CDS ({len(mapa)//3} aa).'}
    pos = [mapa[i0+k] for k in range(3)]
    codon = ''.join(seq[p-1].upper() for p in pos)
    return {
        'erro': None, 'numero_aa': numero_aa,
        'codon': codon, 'aminoacido': CODIGO_GENETICO.get(codon, '???'),
        'pos_inicio': pos[0], 'pos_fim': pos[2], 'pos_gene': pos,
    }

def localizar_stop(seq, coords_cds, proteina_nascente):
    mapa = construir_mapa_cds(seq, coords_cds)
    n_aa = len(proteina_nascente)
    idx  = n_aa * 3
    if idx + 3 > len(mapa):
        p_nt = mapa[-1] if mapa else 0
        codon = seq[p_nt:p_nt+3].upper() if p_nt+3 <= len(seq) else '???'
    else:
        pos = [mapa[idx+k] for k in range(3)]
        codon = ''.join(seq[p-1].upper() for p in pos)
    if codon not in STOP_CODONS:
        return {'erro': None, 'encontrado': False, 'codon': codon,
                'aviso': f'Códon {codon} não é stop codon padrão.'}
    if idx + 3 <= len(mapa):
        pos = [mapa[idx+k] for k in range(3)]
    else:
        pos = [0, 0, 0]
    ultimos = [AA_1_TO_3.get(c.upper(), c) for c in proteina_nascente[-3:]]
    return {
        'erro': None, 'encontrado': True, 'codon': codon,
        'pos_inicio': pos[0], 'pos_fim': pos[2], 'pos_gene': pos,
        'ultimos_aas': ultimos, 'n_aa_total': n_aa,
    }


# ═══════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════

@app.get("/")
def root():
    return {"status": "ExonEditor API online", "version": "1.0.0"}


@app.get("/buscar")
def buscar_gene(termo: str, organismo: str = "Homo sapiens", email: str = EMAIL):
    """Busca registros genômicos no NCBI."""
    try:
        if organismo:
            query = (f'({termo}[Gene Name] OR {termo}[Accession])'
                     f' AND {organismo}[Organism] AND RefSeqGene[Filter]')
        else:
            query = f'({termo}[Gene Name] OR {termo}[Accession]) AND RefSeqGene[Filter]'
        regs = ncbi_buscar(query, email)
        return {"resultados": regs, "total": len(regs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class CarregarPayload(BaseModel):
    acc: str
    email: str = EMAIL
    seq_start: Optional[int] = None
    seq_stop: Optional[int] = None

@app.post("/carregar")
def carregar_registro(payload: CarregarPayload):
    """
    Carrega e parseia um registro GenBank.
    Passo 1: sem seq_start/seq_stop → retorna regiao_sugerida
    Passo 2: com seq_start/seq_stop → retorna sequência recortada com éxons
    """
    try:
        gb  = ncbi_fetch_genbank(payload.acc, payload.email,
                                  payload.seq_start, payload.seq_stop)
        res = parsear_genbank(gb)
        # Remove sequência do payload (muito grande para JSON geral)
        # retorna apenas o necessário para a UI
        return {
            'accession':         res['accession'],
            'definicao':         res['definicao'],
            'organismo':         res['organismo'],
            'total_pb':          len(res['sequencia']),
            'exons':             res['exons'],
            'proteinas':         res['proteinas'],
            'proteina_nascente': res['proteina_nascente'],
            'proteina_madura_seq': res['proteina_madura_seq'],
            'regiao_sugerida':   res['regiao_sugerida'],
            # Sequência é necessária para geração de documentos
            # guardada em sessão pelo front ou enviada de volta nos payloads
            'sequencia':         res['sequencia'],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class AnalisarCodonsPayload(BaseModel):
    sequencia: str
    proteina_nascente: str
    coords_cds: List[List[int]]
    aas_entrada: List[str]

@app.post("/analisar-codons")
def analisar_codons_endpoint(payload: AnalisarCodonsPayload):
    try:
        coords = [tuple(c) for c in payload.coords_cds]
        res = analisar_primeiros_codons(
            payload.sequencia, payload.proteina_nascente,
            coords, payload.aas_entrada)
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class LocalizarCodonPayload(BaseModel):
    sequencia: str
    coords_cds: List[List[int]]
    numero_aa: int

@app.post("/localizar-codon")
def localizar_codon_endpoint(payload: LocalizarCodonPayload):
    try:
        coords = [tuple(c) for c in payload.coords_cds]
        res = localizar_codon_num(payload.sequencia, coords, payload.numero_aa)
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class StopCodonPayload(BaseModel):
    sequencia: str
    coords_cds: List[List[int]]
    proteina_nascente: str

@app.post("/stop-codon")
def stop_codon_endpoint(payload: StopCodonPayload):
    try:
        coords = [tuple(c) for c in payload.coords_cds]
        res = localizar_stop(payload.sequencia, coords, payload.proteina_nascente)
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class GerarDocPayload(BaseModel):
    sequencia: str
    exons: list
    tipo: str = "completo"   # "completo" ou "somente_exons"
    caixa: str = "minuscula"
    fonte: str = "Courier New"
    tamanho: int = 11
    cor_base: str = "#aaaaaa"
    codons_grifar: Optional[list] = None
    utr_regiao: Optional[dict] = None

@app.post("/gerar-documento")
def gerar_documento(payload: GerarDocPayload):
    """Gera documento Word e retorna para download."""
    try:
        cfg = {
            'fonte':   payload.fonte,
            'cor':     payload.cor_base,
            'tamanho': payload.tamanho,
        }
        exons_fmt = []
        for e in payload.exons:
            exons_fmt.append({
                'inicio':  e['inicio'],
                'fim':     e['fim'],
                'numero':  e.get('numero', 1),
                'fonte':   e.get('fonte', payload.fonte),
                'tamanho': e.get('tamanho', payload.tamanho),
                'cor':     e.get('cor', '#000000'),
            })

        if payload.tipo == 'completo':
            buf = gerar_completo_bytes(
                payload.sequencia, cfg, exons_fmt,
                caixa=payload.caixa,
                codons_grifar=payload.codons_grifar,
                utr_regiao=payload.utr_regiao,
            )
        else:
            buf = gerar_so_exons_bytes(
                payload.sequencia, exons_fmt,
                caixa=payload.caixa,
                codons_grifar=payload.codons_grifar,
                utr_regiao=payload.utr_regiao,
            )

        nome = "exon_completo.docx" if payload.tipo == 'completo' else "exon_somente_exons.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nome}"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════
# MODO MANUAL — Upload de .docx
# ═══════════════════════════════════════════════════════
from fastapi import UploadFile, File
import re as _re

def limpar_sequencia_docx(content_bytes: bytes) -> str:
    """Extrai e limpa sequência ATCG de um arquivo .docx."""
    import io
    doc = Document(io.BytesIO(content_bytes))
    texto = ""
    for para in doc.paragraphs:
        t = _re.sub(r'\d+', '', para.text)
        t = _re.sub(r'[\s\r\n\v\f]+', '', t)
        texto += t
    return _re.sub(r'[^ATCGatcg]', '', texto).upper()

def parsear_features_texto(texto: str) -> list:
    """Parseia texto da seção Features do GenBank para extrair éxons."""
    exons = []
    p_pos = _re.compile(r'exon\s+(\d+)\.\.(\d+)', _re.IGNORECASE)
    p_num = _re.compile(r'/number=(\d+)')
    mp = list(p_pos.finditer(texto))
    mn = list(p_num.finditer(texto))
    for i, m in enumerate(mp):
        num = int(mn[i].group(1)) if i < len(mn) else i + 1
        exons.append({'numero': num, 'inicio': int(m.group(1)), 'fim': int(m.group(2))})
    return sorted(exons, key=lambda x: x['numero'])


@app.post("/upload-docx")
async def upload_docx(file: UploadFile = File(...)):
    """Recebe .docx, extrai sequência ATCG limpa e retorna para o front."""
    try:
        content = await file.read()
        seq = limpar_sequencia_docx(content)
        if not seq:
            raise HTTPException(status_code=400, detail="Nenhuma sequência ATCG encontrada no arquivo.")
        return {
            "sequencia": seq,
            "total_pb": len(seq),
            "preview": seq[:120],
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/parsear-features")
async def parsear_features(payload: dict):
    """Parseia texto das Features do GenBank e retorna lista de éxons."""
    texto = payload.get("texto", "")
    exons = parsear_features_texto(texto)
    if not exons:
        raise HTTPException(status_code=400, detail="Nenhum éxon encontrado no texto.")
    return {"exons": exons, "total": len(exons)}
